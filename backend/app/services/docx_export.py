"""Экспорт knowledge blocks в DOCX."""

from io import BytesIO

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_BREAK
from docx.shared import Pt

from app.models.knowledge import KnowledgeBlock


DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _clean_text(value: str | None) -> str:
    return (value or "").strip()


def _add_text_block(document: DocumentObject, text: str) -> None:
    """Сохранить абзацы и пустые строки в человекочитаемом виде."""
    cleaned = _clean_text(text)
    if not cleaned:
        document.add_paragraph("Материал пуст.")
        return

    for line in cleaned.splitlines():
        if line.strip():
            document.add_paragraph(line.rstrip())
        else:
            document.add_paragraph()


def _add_material_section(
    document: DocumentObject,
    title: str,
    text: str | None,
) -> None:
    document.add_heading(title, level=2)
    _add_text_block(document, _clean_text(text))


def _add_block(document: DocumentObject, block: KnowledgeBlock) -> None:
    title = _clean_text(block.video_title) or f"Материал {block.video_job_id}"
    document.add_heading(title, level=1)
    _add_material_section(document, "Конспект", block.summary_text)
    _add_material_section(document, "Методичка", block.manual_text)
    _add_material_section(document, "Чек-лист", block.checklist_text)


def _configure_document(document: DocumentObject) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)


def render_single_block_docx(block: KnowledgeBlock) -> bytes:
    """Собрать DOCX для одного knowledge block."""
    document = Document()
    _configure_document(document)
    _add_block(document, block)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def render_master_docx(blocks: list[KnowledgeBlock]) -> bytes:
    """Собрать master-DOCX из блоков в порядке добавления."""
    document = Document()
    _configure_document(document)
    document.add_heading("Master-документ базы знаний", level=0)

    for index, block in enumerate(blocks):
        if index > 0:
            paragraph = document.add_paragraph()
            paragraph.add_run().add_break(WD_BREAK.PAGE)
        _add_block(document, block)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
